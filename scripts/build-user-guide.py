from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

OUT=Path(r"C:\nisithrportal\artifacts\hr-portal\public")
CLEAN_JOB=Path(r"C:\Users\Mukombo\AppData\Local\Temp\codex-clipboard-aa44cbeb-44d5-42be-bfc1-3e421b343723.png")
CLEAN_ORG=Path(r"C:\Users\Mukombo\AppData\Local\Temp\codex-clipboard-55ac8661-744a-48bd-af27-e922b7be5182.png")
NAVY=RGBColor(15,42,77); RED=RGBColor(210,22,55); GREY=RGBColor(91,99,110)

def shade(cell,color):
 p=cell._tc.get_or_add_tcPr(); e=OxmlElement('w:shd');e.set(qn('w:fill'),color);p.append(e)
def base(title,audience):
 d=Document(); sec=d.sections[0]; sec.top_margin=Inches(.7);sec.bottom_margin=Inches(.7);sec.left_margin=sec.right_margin=Inches(.8)
 normal=d.styles['Normal'];normal.font.name='Aptos';normal.font.size=Pt(10.5);normal.font.color.rgb=RGBColor(45,48,53);normal.paragraph_format.space_after=Pt(6)
 for name,size,color,before,after in [('Heading 1',18,NAVY,18,7),('Heading 2',13,NAVY,12,5),('Heading 3',11,RED,9,4)]:
  s=d.styles[name];s.font.name='Aptos Display';s.font.size=Pt(size);s.font.bold=True;s.font.color.rgb=color;s.paragraph_format.space_before=Pt(before);s.paragraph_format.space_after=Pt(after)
 p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=p.add_run('PNG NISIT HR PORTAL');r.bold=True;r.font.size=Pt(12);r.font.color.rgb=RED
 p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=p.add_run(title);r.bold=True;r.font.size=Pt(26);r.font.color.rgb=NAVY
 p=d.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=p.add_run(audience);r.italic=True;r.font.size=Pt(12);r.font.color.rgb=GREY
 note(d,'Read this first','This is a practical task guide. Complete one step at a time. Names of buttons and menus match the portal. Screenshots in this edition are clean examples only; they do not show error messages.')
 return d
def note(d,label,text):
 t=d.add_table(rows=1,cols=1);t.alignment=WD_TABLE_ALIGNMENT.CENTER;c=t.cell(0,0);shade(c,'EEF4FB');c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER;p=c.paragraphs[0];r=p.add_run(label+'  ');r.bold=True;r.font.color.rgb=NAVY;p.add_run(text)
def h(d,x):d.add_paragraph(x,'Heading 1')
def h2(d,x):d.add_paragraph(x,'Heading 2')
def intro(d,x):p=d.add_paragraph(x);p.paragraph_format.space_after=Pt(8)
def numbered(d,arr):
 for x in arr:d.add_paragraph(x,style='List Number')
def bullets(d,arr):
 for x in arr:d.add_paragraph(x,style='List Bullet')
def matrix(d,headers,rows):
 t=d.add_table(rows=1,cols=len(headers));t.style='Table Grid';t.alignment=WD_TABLE_ALIGNMENT.CENTER
 for c,x in zip(t.rows[0].cells,headers):c.text=x;shade(c,'DCE9F7');c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
 for row in rows:
  cells=t.add_row().cells
  for c,x in zip(cells,row):c.text=x;c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
def picture(d,path,caption):
 if path.exists():
  d.add_picture(str(path),width=Inches(6.4));p=d.add_paragraph('Figure: '+caption);p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.runs[0].italic=True;p.runs[0].font.size=Pt(9);p.runs[0].font.color.rgb=GREY
def end(d,label):
 p=d.sections[0].footer.paragraphs[0];p.alignment=WD_ALIGN_PARAGRAPH.CENTER;r=p.add_run('PNG NISIT HR Portal | '+label);r.font.size=Pt(8);r.font.color.rgb=GREY

# Applicant / non-staff guide
d=base('Applicant and Non-Staff Guide','From registration to application tracking')
h(d,'Guide map');matrix(d,['Use this guide when you need to…','Open this module'],[('Create or regain access','Sign in / Forgot password'),('Search and save a job','Job Vacancies'),('Start, save or submit an application','Job vacancy > Apply Now'),('Check progress after submission','My Applications'),('Change account details','My Account')])
h(d,'1. Create your account and protect it');intro(d,'Use one personal email address. It is how the portal identifies your applications and sends updates.')
numbered(d,['Open the portal and select Register as applicant.','Enter your legal name as it appears on your documents. Enter an email address you can access.','Create a strong password and complete registration.','Sign in. Check the top-right welcome area: it should show your name and email.','If you cannot sign in, select Forgot password? Do not repeatedly guess your password.'])
note(d,'Good practice','Never share your password. On a shared computer, always select Sign out when you finish.')
h(d,'2. Find a suitable vacancy');numbered(d,['Select Job Vacancies in the left menu.','Use search to enter a job title, location or keyword. Use filters where available.','Open a vacancy by selecting its title.','Read the overview, responsibilities, qualifications, location, employment type, salary information (if shown), required documents and closing date.','Only apply when the vacancy is Open or Published and the closing date has not passed.'])
picture(d,CLEAN_JOB,'A clean job detail page. Select Apply Now after reviewing the vacancy.')
h(d,'3. Apply - complete every section');intro(d,'The wizard can be completed in more than one sitting. Use Save Draft after meaningful progress. A draft is not submitted and is not visible to recruiters as an application.')
matrix(d,['Wizard section','What to do before moving on'],[('Personal','Use your legal names and correct personal details.'),('Contact','Confirm your email and phone number; include a current address where requested.'),('Availability','State when and where you can work honestly.'),('Education and experience','Add relevant records; use real dates and institution/employer names.'),('Skills','List skills that you can demonstrate.'),('Documents','Upload your CV and the supporting documents requested for that vacancy.'),('Screening','Answer every required question clearly and truthfully.'),('Declarations','Read each declaration before confirming it.'),('Diversity','Optional. Leaving this section blank does not block submission.')])
h2(d,'Upload your CV and evidence');numbered(d,['Open the Documents step.','Select Click to upload CV. Choose PDF, DOC or DOCX, up to 10 MB.','Wait until the document has finished uploading before selecting Next.','For certificates, choose a document type under Additional Documents, then upload the matching file.','Check the displayed file name before continuing.'])
note(d,'If upload does not complete','Keep the wizard open. Check that the file is PDF/DOC/DOCX and below 10 MB. Retry once after confirming your internet connection. If it still fails, contact support with the file type, size and a screenshot.')
h(d,'4. Submit and track');numbered(d,['At the final step, review the summary and declarations.','Select Submit Application once. Wait for confirmation; do not refresh or click repeatedly.','Open My Applications to see the application status and timeline.','Read portal notifications and email for requests, interview invitations or updates.','Update My Account if your email or phone number changes.'])
h(d,'5. Common questions');matrix(d,['Message or situation','What it means / what to do'],[('Job is not accepting applications','The vacancy is closed, withdrawn or no longer public. Select another vacancy.'),('You already applied','Open My Applications instead of creating a duplicate.'),('Document upload failed','Check type, size and connection. Keep a copy of the file; contact support if it continues.'),('Forgot password','Use Forgot password? and follow the reset email.'),('No update yet','Recruitment stages take time. Check My Applications and notifications regularly.')])
h(d,'6. Before you leave');bullets(d,['Confirm the application appears in My Applications.','Keep original documents ready for verification if requested.','Do not send documents to unofficial email addresses.','Sign out if you are using a shared device.']);end(d,'Applicant and Non-Staff Guide');d.save(OUT/'NISIT-HR-Portal-Applicant-User-Guide.docx')

# Staff guide
s=base('Staff User Guide','Employee self-service, approvals and workforce workflows')
h(s,'Guide map');matrix(s,['Routine task','Module'],[('Update personal account / notification preferences','My Account'),('Apply for leave or review leave','Leave & Absence'),('Clock attendance','Attendance Clock'),('Access authorised records','Document Vault / HR Letters'),('Manage people and contracts (per role)','Employees / Contracts / Onboarding / Offboarding'),('Review recruitment','Job Vacancies / Applications / Recruitment Workflow'),('Export information','Any portal table with Export')])
h(s,'1. Start each day');numbered(s,['Sign in with your staff account.','Check the welcome panel for your full name, email, date and live time.','Open the notification bell and clear or act on urgent items.','Use the grouped navigation menu; expand only the area you need.','Select My Account to verify your contact details and profile settings.'])
h(s,'2. Use tables correctly');intro(s,'The portal uses the same table controls in operational modules. They help you find records without losing your current view.')
numbered(s,['Use Search for a quick keyword search.','Select Filters for a field-specific search, then Clear filters when done.','Select a column heading to sort ascending; select again to reverse it.','Select Columns to hide, show or reset columns. Resize a column by dragging its right edge.','Use the checkbox column to select records. Use Actions only when your role permits the chosen action.','Use Export to download CSV, TSV or JSON. Export only data you are authorised to handle.','Use Rows per page, the page box and navigation arrows for long lists.'])
h(s,'3. Employee self-service');h2(s,'Leave and absence');numbered(s,['Open Leave & Absence.','Select Apply for Leave.','Choose the correct leave category.','Enter a start date and an end date. The end date cannot be before the start date.','Enter a concise reason and handover details.','Select Submit Application, then monitor the status until it is approved, returned or rejected.'])
h2(s,'Attendance, documents and letters');bullets(s,['Use Attendance Clock only for your own work time unless you have a specific authorised role.','Use Document Vault to view documents you are allowed to access. Use View Document in the portal rather than relying on broken external links.','Use HR Letters for available letters. Keep downloaded files secure.'])
h(s,'4. Supervisor, HR and manager workflows');matrix(s,['Workflow','Minimum safe sequence'],[('Recruitment','Create/publish vacancy; review applications; record assessment/interview outcomes; move candidates only through valid stages; keep notes factual.'),('Employee records','Verify identity and contact details; use an employable date of birth; set department, position, supervisor, appointment date and status correctly.'),('Onboarding','Create workflow; assign owners and dates; complete tasks and upload required evidence; monitor outstanding items.'),('Contracts','Use dates only for time-bound contracts; record renewals before expiry; upload or replace signed versions through the portal.'),('Offboarding','Initiate separation; assign clearance tasks; complete assets, access, final pay and records steps before closure.'),('Approvals','Review the request, documents and dates; approve, return or reject with a clear comment; honour approved delegation periods.')])
h(s,'5. Organisation, dashboard and reports');numbered(s,['Open Dashboard and select a card/chart to drill into the underlying records.','Open Org Hierarchy to view filled and vacant positions. Use the filters before interpreting headcount.','Open Standard Reports; choose the report, apply appropriate dates and filters, then export or print.','Do not treat empty or non-applicable dates as a renewal/expiry. Confirm whether the record is a time-bound contract first.'])
picture(s,CLEAN_ORG,'Organisation hierarchy with filled/vacant cards and filters.')
h(s,'6. Data protection and error handling');bullets(s,['Only access records that your role allows.','Use accurate, real dates. Do not use placeholder dates for birth, appointment, renewal or separation fields.','Do not email exports or personal documents to unauthorised recipients.','Read error messages before retrying. Save your work or take a screenshot before refreshing.','Report a fault with the module name, record reference, time, steps taken and a screenshot.'])
h(s,'7. End-of-task checklist');bullets(s,['The record status is correct.','All applicable dates are logical and complete.','Required evidence/document uploads show as complete.','Any approval comment is clear and professional.','Sensitive information is not left open on a shared computer.']);end(s,'Staff User Guide');s.save(OUT/'NISIT-HR-Portal-Staff-User-Guide.docx')
print('Rebuilt applicant and staff guides')
